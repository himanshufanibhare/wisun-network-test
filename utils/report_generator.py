#!/usr/bin/env python3
"""
Report Generator Utilities
Functions to generate Wi-SUN tree reports in different formats (PDF, Word, TXT)
"""

import os
import json
from datetime import datetime
from io import BytesIO
import tempfile

def parse_wisun_tree_data(tree_output):
    """
    Parse Wi-SUN tree output to extract structured data
    Returns: dict with parsed network information
    """
    if not tree_output:
        return {}
    
    lines = tree_output.strip().split('\n')
    network_info = {}
    devices = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Parse network configuration
        if line.startswith('network_name:'):
            network_info['network_name'] = line.split(':', 1)[1].strip()
        elif line.startswith('fan_version:'):
            network_info['fan_version'] = line.split(':', 1)[1].strip()
        elif line.startswith('domain:'):
            network_info['domain'] = line.split(':', 1)[1].strip()
        elif line.startswith('phy_mode_id:'):
            network_info['phy_mode_id'] = line.split(':', 1)[1].strip()
        elif line.startswith('chan_plan_id:'):
            network_info['chan_plan_id'] = line.split(':', 1)[1].strip()
        elif line.startswith('panid:'):
            network_info['panid'] = line.split(':', 1)[1].strip()
        elif line.startswith('size:'):
            network_info['size'] = line.split(':', 1)[1].strip()
        # Parse device entries (IPv6 addresses)
        elif '::' in line:
            # Extract IPv6 address
            if line.startswith('├─') or line.startswith('└─'):
                # This is a device in the tree
                ipv6 = line.split()[-1] if line.split() else ''
                if '::' in ipv6:
                    devices.append({
                        'ipv6': ipv6,
                        'indent_level': len(line) - len(line.lstrip()) // 2,
                        'line': line
                    })
            elif line.startswith('fd12:') or line.startswith('FD12:'):
                # Direct IPv6 listing
                devices.append({
                    'ipv6': line,
                    'indent_level': 0,
                    'line': line
                })
    
    return {
        'network_info': network_info,
        'devices': devices,
        'total_devices': len(devices)
    }

def generate_txt_report(tree_output, device_count, timestamp):
    """
    Generate TXT format report
    """
    parsed_data = parse_wisun_tree_data(tree_output)
    
    report = []
    report.append("=" * 60)
    report.append("           Wi-SUN Network Tree Status Report")
    report.append("=" * 60)
    report.append("")
    report.append(f"Generated: {timestamp}")
    report.append(f"Total Wi-SUN Devices: {device_count}")
    report.append("")
    
    # Network Information
    if parsed_data.get('network_info'):
        report.append("Network Configuration:")
        report.append("-" * 25)
        for key, value in parsed_data['network_info'].items():
            report.append(f"{key.replace('_', ' ').title()}: {value}")
        report.append("")
    
    # Device List
    report.append("Connected Devices:")
    report.append("-" * 18)
    report.append("")
    
    if tree_output:
        # Use original tree output for better formatting
        report.extend(tree_output.split('\n'))
    else:
        report.append("No device data available")
    
    report.append("")
    report.append("=" * 60)
    report.append("End of Report")
    report.append("=" * 60)
    
    return '\n'.join(report)

def generate_pdf_report(tree_output, device_count, timestamp):
    """
    Generate PDF format report using reportlab
    """
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, KeepTogether
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=1*inch)
        
        # Get styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER,
            textColor=colors.darkblue
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            textColor=colors.darkgreen
        )
        
        # Build content
        story = []
        
        # Title
        story.append(Paragraph("Wi-SUN Network Tree Status Report", title_style))
        story.append(Spacer(1, 20))
        
        # Summary table
        summary_data = [
            ['Generated:', timestamp],
            ['Total Wi-SUN Devices:', str(device_count)],
        ]
        
        summary_table = Table(summary_data, colWidths=[2*inch, 3*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(summary_table)
        story.append(Spacer(1, 30))
        
        # Network Tree
        story.append(Paragraph("Network Tree Structure", heading_style))
        
        if tree_output:
            # Format tree output for PDF
            tree_lines = tree_output.split('\n')
            tree_data = [[line] for line in tree_lines if line.strip()]
            
            if tree_data:
                tree_table = Table(tree_data, colWidths=[6*inch])
                tree_table.setStyle(TableStyle([
                    ('FONTNAME', (0, 0), (-1, -1), 'Courier'),
                    ('FONTSIZE', (0, 0), (-1, -1), 8),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 6),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 6),
                    ('TOPPADDING', (0, 0), (-1, -1), 3),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
                ]))
                story.append(tree_table)
        else:
            story.append(Paragraph("No network tree data available", styles['Normal']))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
        
    except ImportError:
        raise ImportError("reportlab library is required for PDF generation")

def generate_word_report(tree_output, device_count, timestamp):
    """
    Generate Word format report using python-docx
    """
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import RGBColor
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Wi-SUN Network Tree Status Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary information
        doc.add_heading('Report Summary', level=1)
        
        summary_table = doc.add_table(rows=2, cols=2)
        summary_table.style = 'Table Grid'
        
        # Header row
        hdr_cells = summary_table.rows[0].cells
        hdr_cells[0].text = 'Generated'
        hdr_cells[1].text = timestamp
        
        hdr_cells = summary_table.rows[1].cells
        hdr_cells[0].text = 'Total Wi-SUN Devices'
        hdr_cells[1].text = str(device_count)
        
        # Make header bold
        for cell in summary_table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Network Tree Section
        doc.add_heading('Network Tree Structure', level=1)
        
        if tree_output:
            # Add tree output as preformatted text
            tree_paragraph = doc.add_paragraph()
            tree_run = tree_paragraph.add_run(tree_output)
            tree_run.font.name = 'Courier New'
            tree_run.font.size = 10
        else:
            doc.add_paragraph('No network tree data available')
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
        
    except ImportError:
        raise ImportError("python-docx library is required for Word document generation")

def generate_json_report(tree_output, device_count, timestamp):
    """
    Generate JSON format report
    """
    parsed_data = parse_wisun_tree_data(tree_output)
    
    report = {
        "report_info": {
            "title": "Wi-SUN Network Tree Status Report",
            "generated": timestamp,
            "total_wisun_devices": device_count
        },
        "network_config": parsed_data.get('network_info', {}),
        "devices": parsed_data.get('devices', []),
        "raw_tree_output": tree_output
    }
    
    return json.dumps(report, indent=2)

def generate_csv_report(tree_output, device_count, timestamp):
    """
    Generate CSV format report
    """
    import csv
    from io import StringIO
    
    parsed_data = parse_wisun_tree_data(tree_output)
    
    output = StringIO()
    writer = csv.writer(output)
    
    # Header
    writer.writerow(['Wi-SUN Network Tree Status Report'])
    writer.writerow(['Generated', timestamp])
    writer.writerow(['Total Wi-SUN Devices', device_count])
    writer.writerow([])
    
    # Network Configuration
    if parsed_data.get('network_info'):
        writer.writerow(['Network Configuration'])
        for key, value in parsed_data['network_info'].items():
            writer.writerow([key.replace('_', ' ').title(), value])
        writer.writerow([])
    
    # Device List
    writer.writerow(['Device Information'])
    writer.writerow(['IPv6 Address', 'Indent Level', 'Raw Line'])
    
    for device in parsed_data.get('devices', []):
        writer.writerow([device.get('ipv6', ''), device.get('indent_level', ''), device.get('line', '')])
    
    return output.getvalue()

def generate_xml_report(tree_output, device_count, timestamp):
    """
    Generate XML format report
    """
    parsed_data = parse_wisun_tree_data(tree_output)
    
    xml_lines = []
    xml_lines.append('<?xml version="1.0" encoding="UTF-8"?>')
    xml_lines.append('<wisun_network_report>')
    xml_lines.append('  <report_info>')
    xml_lines.append('    <title>Wi-SUN Network Tree Status Report</title>')
    xml_lines.append(f'    <generated>{timestamp}</generated>')
    xml_lines.append(f'    <total_wisun_devices>{device_count}</total_wisun_devices>')
    xml_lines.append('  </report_info>')
    
    # Network Configuration
    if parsed_data.get('network_info'):
        xml_lines.append('  <network_config>')
        for key, value in parsed_data['network_info'].items():
            xml_lines.append(f'    <{key}>{value}</{key}>')
        xml_lines.append('  </network_config>')
    
    # Devices
    xml_lines.append('  <devices>')
    for device in parsed_data.get('devices', []):
        xml_lines.append('    <device>')
        xml_lines.append(f'      <ipv6>{device.get("ipv6", "")}</ipv6>')
        xml_lines.append(f'      <indent_level>{device.get("indent_level", "")}</indent_level>')
        xml_lines.append(f'      <raw_line><![CDATA[{device.get("line", "")}]]></raw_line>')
        xml_lines.append('    </device>')
    xml_lines.append('  </devices>')
    
    # Raw tree output
    xml_lines.append('  <raw_tree_output>')
    xml_lines.append(f'    <![CDATA[{tree_output}]]>')
    xml_lines.append('  </raw_tree_output>')
    
    xml_lines.append('</wisun_network_report>')
    
def get_file_extension(format_type):
    """Get file extension for format type"""
    extensions = {
        'txt': '.txt',
        'pdf': '.pdf',
        'word': '.docx',
        'json': '.json',
        'csv': '.csv',
        'xml': '.xml'
    }
    return extensions.get(format_type, '.txt')

def get_mimetype(format_type):
    """Get MIME type for format type"""
    mimetypes = {
        'txt': 'text/plain',
        'pdf': 'application/pdf',
        'word': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'json': 'application/json',
        'csv': 'text/csv',
        'xml': 'application/xml'
    }
    return mimetypes.get(format_type, 'text/plain')

def generate_filename(format_type, timestamp=None):
    """Generate filename for download"""
    if not timestamp:
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    else:
        # Convert timestamp to filename-safe format
        timestamp = timestamp.replace(':', '').replace('-', '').replace(' ', '_')
    
    extension = get_file_extension(format_type)
    return f"wisun_tree_report_{timestamp}{extension}"