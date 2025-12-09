"""
Test Result Writer Module
Handles writing test results in different formats (TXT, PDF, Word)
"""

import os
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches as DocxInches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class TestResultWriter:
    def __init__(self, test_type, output_format, timestamp=None):
        print(f"DEBUG: TestResultWriter.__init__ called with test_type='{test_type}', output_format='{output_format}'")
        self.test_type = test_type
        self.output_format = output_format.lower()
        print(f"DEBUG: Normalized output_format to: '{self.output_format}'")
        self.timestamp = timestamp or datetime.now().strftime("%Y%m%d_%H%M%S")
        self.base_dir = "reports"
        
        # Create directories if they don't exist
        self.output_dir = os.path.join(self.base_dir, self.output_format)
        os.makedirs(self.output_dir, exist_ok=True)
        print(f"DEBUG: Created output directory: {self.output_dir}")
        
        # Set up file paths
        self.filename = f"{test_type}_test_{self.timestamp}"
        if self.output_format == 'txt':
            self.file_path = os.path.join(self.output_dir, f"{self.filename}.txt")
        elif self.output_format == 'pdf':
            self.file_path = os.path.join(self.output_dir, f"{self.filename}.pdf")
        elif self.output_format == 'word':
            self.file_path = os.path.join(self.output_dir, f"{self.filename}.docx")
        
        print(f"DEBUG: File path set to: {self.file_path}")
        
        # Store all results for table format
        self.results = []
        
        # Initialize the file
        print(f"DEBUG: Initializing file for output_format: {self.output_format}")
        self._initialize_file()
    
    def _initialize_file(self):
        """Initialize the output file with headers"""
        if self.output_format == 'txt':
            self._initialize_txt()
        elif self.output_format == 'pdf':
            self._initialize_pdf()
        elif self.output_format == 'word':
            self._initialize_word()
    
    def _initialize_txt(self):
        """Initialize TXT file"""
        with open(self.file_path, 'w', encoding='utf-8') as f:
            f.write(f"Wi-SUN Network Test Report\n")
            f.write(f"{'=' * 50}\n")
            f.write(f"Test Type: {self.test_type.upper()}\n")
            f.write(f"Started: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"{'=' * 50}\n\n")
    
    def _initialize_pdf(self):
        """Initialize PDF file"""
        self.pdf_story = []
        self.pdf_styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.pdf_styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        # Add title and header info
        self.pdf_story.append(Paragraph("Wi-SUN Network Test Report", title_style))
        self.pdf_story.append(Spacer(1, 12))
        
        header_data = [
            ['Test Type:', self.test_type.upper()],
            ['Started:', datetime.now().strftime('%Y-%m-%d %H:%M:%S')],
            ['Format:', 'PDF Report']
        ]
        
        header_table = Table(header_data, colWidths=[2*inch, 4*inch])
        header_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.grey),
            ('TEXTCOLOR', (0, 0), (0, -1), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 12),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('BACKGROUND', (1, 0), (1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        self.pdf_story.append(header_table)
        self.pdf_story.append(Spacer(1, 20))
    
    def _initialize_word(self):
        """Initialize Word document"""
        self.doc = Document()
        
        # Add title
        title = self.doc.add_heading('Wi-SUN Network Test Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add header info
        self.doc.add_heading('Test Information', level=1)
        info_table = self.doc.add_table(rows=3, cols=2)
        info_table.style = 'Table Grid'
        
        cells = info_table.rows[0].cells
        cells[0].text = 'Test Type:'
        cells[1].text = self.test_type.upper()
        
        cells = info_table.rows[1].cells
        cells[0].text = 'Started:'
        cells[1].text = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        cells = info_table.rows[2].cells
        cells[0].text = 'Format:'
        cells[1].text = 'Word Document'
        
        self.doc.add_paragraph('')  # Add spacing
        self.doc.add_heading('Test Results', level=1)
    
    def append_result(self, device_result):
        """Store device result for table format"""
        print(f"DEBUG: append_result called with device_result: {device_result}")
        # Add serial number
        device_result['sr_no'] = len(self.results) + 1
        self.results.append(device_result)
        print(f"DEBUG: Result appended. Total results now: {len(self.results)}")
    
    def _get_table_headers(self):
        """Get table headers based on test type"""
        base_headers = ['Sr No.', 'IP Address', 'Device Label', 'Hop Count']
        
        if self.test_type == 'ping':
            return base_headers + ['Packets TX', 'Packets RX', 'Loss (%)', 'Min RTT (ms)', 'Max RTT (ms)', 'Avg RTT (ms)', 'Mdev (ms)', 'Connection Status']
        elif self.test_type == 'rssi' or self.test_type == 'rssl':
            return base_headers + ['RSL In (dBm)', 'RSL Out (dBm)', 'Connection Status']
        elif self.test_type == 'rpl':
            return base_headers + ['RPL Rank', 'Connection Status']
        elif self.test_type == 'disconnections':
            return base_headers + ['Disconnected Total', 'Connection Status']
        elif self.test_type == 'availability':
            return base_headers + ['Availability Status', 'Connection Status']
        else:
            return base_headers + ['Connection Status']
    
    def _get_table_row(self, result):
        """Generate a table row for the given result"""
        # Find the position of this result in the results list
        sr_no = str(self.results.index(result) + 1)
        ip = result.get('ip', 'N/A')
        device_label = result.get('device_label', result.get('label', 'Unknown'))  # Check both field names
        hop_count = str(result.get('hop_count', 'N/A'))
        
        # Base row with common fields
        row = [sr_no, ip, device_label, hop_count]
        
        if self.test_type == 'ping':
            packets_tx = str(result.get('packets_tx', 'N/A'))
            packets_rx = str(result.get('packets_rx', 'N/A'))
            loss_percent = f"{result.get('loss_percent', 'N/A')}%" if result.get('loss_percent') is not None and result.get('loss_percent') != '-' else (result.get('loss_percent', 'N/A') if result.get('loss_percent') == '-' else 'N/A')
            min_rtt = str(result.get('min_time', 'N/A'))  # Changed from min_rtt to min_time
            max_rtt = str(result.get('max_time', 'N/A'))  # Changed from max_rtt to max_time
            avg_rtt = str(result.get('avg_time', 'N/A'))  # Changed from avg_rtt to avg_time
            mdev = str(result.get('mdev_time', 'N/A'))    # Changed from mdev to mdev_time
            connection_status = result.get('connection_status', 'Unknown')
            
            row.extend([packets_tx, packets_rx, loss_percent, min_rtt, max_rtt, avg_rtt, mdev, connection_status])
            
        elif self.test_type == 'rssi' or self.test_type == 'rssl':
            rsl_in = str(result.get('rsl_in', 'N/A'))
            rsl_out = str(result.get('rsl_out', 'N/A'))
            connection_status = result.get('connection_status', 'Unknown')
            
            row.extend([rsl_in, rsl_out, connection_status])
            
        elif self.test_type == 'rpl':
            rpl_rank = str(result.get('rpl_data', 'N/A'))
            connection_status = result.get('connection_status', 'Unknown')
            
            row.extend([rpl_rank, connection_status])
            
        elif self.test_type == 'disconnections':
            # Use disconnected_total field from the test data
            disconnected_total = str(result.get('disconnected_total', 'N/A'))
            connection_status = result.get('connection_status', 'Unknown')
            
            row.extend([disconnected_total, connection_status])
            
        elif self.test_type == 'availability':
            ap = result.get('availability_percent', 'N/A')
            if isinstance(ap, (int, float)):
                availability_percent = f"{ap}%"
            elif ap == "No response or CoAP error":
                availability_percent = ap
            else:
                availability_percent = str(ap)
            connection_status = 'AVAILABLE' if isinstance(ap, (int, float)) and ap > 90 else 'UNAVAILABLE'
            row.extend([availability_percent, connection_status])
        
        return row
    
    def append_summary(self, summary_text):
        """Store summary text for finalization"""
        self.summary_text = summary_text
    
    def write_summary(self, summary_text):
        """Store summary text for finalization (alias for append_summary)"""
        print(f"DEBUG: write_summary called with: {summary_text}")
        self.append_summary(summary_text)
    
    def add_wisun_tree(self, tree_output, timestamp):
        """Add Wi-SUN tree output to the report"""
        print(f"DEBUG: add_wisun_tree called with tree output length: {len(tree_output)}")
        self.wisun_tree_output = tree_output
        self.wisun_tree_timestamp = timestamp
    
    def finalize(self):
        """Finalize and save the file with table format"""
        print(f"DEBUG: TestResultWriter.finalize() called for output_format: {self.output_format}")
        print(f"DEBUG: Number of results to write: {len(self.results)}")
        
        if self.output_format == 'txt':
            self._generate_txt_table()
        elif self.output_format == 'pdf':
            self._generate_pdf_table()
        elif self.output_format == 'word':
            print(f"DEBUG: Calling _generate_word_table()")
            self._generate_word_table()
        
        print(f"DEBUG: Finalization complete. File path: {self.file_path}")
        print(f"DEBUG: File exists after finalization: {os.path.exists(self.file_path)}")
        if os.path.exists(self.file_path):
            file_size = os.path.getsize(self.file_path)
            print(f"DEBUG: File size: {file_size} bytes")
        
        return self.file_path
    
    def _generate_txt_table(self):
        """Generate TXT file with table format"""
        headers = self._get_table_headers()
        
        with open(self.file_path, 'a', encoding='utf-8') as f:
            f.write("Test Results\n")
            f.write("=" * 120 + "\n")
            
            # Calculate column widths
            col_widths = []
            for i, header in enumerate(headers):
                max_width = len(header)
                for result in self.results:
                    row_data = self._get_table_row(result)
                    if i < len(row_data):
                        max_width = max(max_width, len(str(row_data[i])))
                col_widths.append(max_width + 2)  # Add padding
            
            # Create table border
            border_line = "+" + "+".join("-" * width for width in col_widths) + "+"
            f.write(border_line + "\n")
            
            # Write header row
            header_row = "|"
            for i, header in enumerate(headers):
                header_row += f" {header:<{col_widths[i]-1}}|"
            f.write(header_row + "\n")
            f.write(border_line + "\n")
            
            # Write data rows
            for result in self.results:
                row_data = self._get_table_row(result)
                data_row = "|"
                for i, data in enumerate(row_data):
                    if i < len(col_widths):
                        data_row += f" {str(data):<{col_widths[i]-1}}|"
                f.write(data_row + "\n")
            
            f.write(border_line + "\n")
            
            # Add summary
            if hasattr(self, 'summary_text'):
                f.write(f"\nTEST SUMMARY\n")
                f.write(f"{'=' * 80}\n")
                f.write(f"{self.summary_text}\n")
                f.write(f"Completed: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")

    def _generate_pdf_table(self):
        """Generate PDF file with table format"""
        headers = self._get_table_headers()
        
        # Add table title
        self.pdf_story.append(Paragraph("Test Results", 
                                      ParagraphStyle('TableTitle', 
                                                   parent=self.pdf_styles['Heading2'],
                                                   fontSize=16, spaceAfter=20)))
        
        # Prepare table data
        table_data = [headers]  # Header row
        for result in self.results:
            table_data.append(self._get_table_row(result))
        
        # Calculate column widths
        num_cols = len(headers)
        col_width = 7.5 * inch / num_cols  # Distribute evenly across page width
        col_widths = [col_width] * num_cols
        
        # Create table
        table = Table(table_data, colWidths=col_widths, repeatRows=1)
        table.setStyle(TableStyle([
            # Header styling
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 8),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            
            # Data styling
            ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 1), (-1, -1), 7),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            
            # Alternate row colors
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey])
        ]))
        
        self.pdf_story.append(table)
        
        # Add summary
        if hasattr(self, 'summary_text'):
            self.pdf_story.append(Spacer(1, 20))
            summary_style = ParagraphStyle(
                'Summary',
                parent=self.pdf_styles['Heading2'],
                fontSize=14,
                spaceAfter=12
            )
            self.pdf_story.append(Paragraph("Test Summary", summary_style))
            self.pdf_story.append(Paragraph(self.summary_text, self.pdf_styles['Normal']))
            self.pdf_story.append(Paragraph(f"Completed: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", 
                                          self.pdf_styles['Normal']))
        
        # Build PDF
        doc = SimpleDocTemplate(self.file_path, pagesize=letter)
        doc.build(self.pdf_story)
    
    def _generate_word_table(self):
        """Generate Word document with table format"""
        print(f"DEBUG: _generate_word_table() called")
        headers = self._get_table_headers()
        print(f"DEBUG: Table headers: {headers}")
        
        # Add table title
        self.doc.add_heading('Test Results', level=1)
        
        # Create table
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        print(f"DEBUG: Created table with {len(headers)} columns")
        
        # Set column widths based on content
        if self.test_type == 'ping':
            # Ping test has 12 columns, adjust to fit on page
            col_widths = [DocxInches(0.4), DocxInches(1.8), DocxInches(1.0), DocxInches(0.6), 
                         DocxInches(0.6), DocxInches(0.6), DocxInches(0.6), DocxInches(0.6), 
                         DocxInches(0.6), DocxInches(0.6), DocxInches(0.6), DocxInches(1.0)]
        else:
            # Other tests have fewer columns
            col_widths = [DocxInches(0.5), DocxInches(2.5), DocxInches(1.5), DocxInches(1.0), 
                         DocxInches(1.5), DocxInches(1.5)]
        
        # Apply column widths
        for i, width in enumerate(col_widths[:len(headers)]):
            for cell in table.columns[i].cells:
                cell.width = width
        
        # Set header row
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Make header bold and center-aligned
            for paragraph in header_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
        
        # Add data rows
        for result in self.results:
            row_data = self._get_table_row(result)
            row_cells = table.add_row().cells
            for i, data in enumerate(row_data):
                row_cells[i].text = str(data)
                # Center-align numeric data columns
                if i > 3:  # After basic info columns (Sr No, IP, Label, Hop Count)
                    for paragraph in row_cells[i].paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Set table alignment
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add summary if available
        if hasattr(self, 'summary_text'):
            self.doc.add_paragraph('')  # Add spacing
            self.doc.add_heading('Test Summary', level=2)
            self.doc.add_paragraph(self.summary_text)
            self.doc.add_paragraph(f"Completed: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Add Wi-SUN tree at the end if available
        if hasattr(self, 'wisun_tree_output'):
            self.doc.add_page_break()
            self.doc.add_heading('Wi-SUN Network Tree', level=1)
            self.doc.add_paragraph(f"Generated: {self.wisun_tree_timestamp}")
            self.doc.add_paragraph('')  # Add spacing
            
            # Add tree output with monospace font
            tree_paragraph = self.doc.add_paragraph()
            tree_run = tree_paragraph.add_run(self.wisun_tree_output)
            tree_run.font.name = 'Courier New'
            tree_run.font.size = DocxInches(0.1)  # Readable font size
        
        # Save document
        print(f"DEBUG: Saving Word document to: {self.file_path}")
        self.doc.save(self.file_path)
        print(f"DEBUG: Word document saved successfully")
        
        # Verify file was created
        if os.path.exists(self.file_path):
            file_size = os.path.getsize(self.file_path)
            print(f"DEBUG: Verified file exists with size: {file_size} bytes")
        else:
            print(f"ERROR: Word file was not created at: {self.file_path}")
    
    def get_file_path(self):
        """Get the current file path"""
        return self.file_path