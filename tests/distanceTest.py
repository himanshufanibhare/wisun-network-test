import re
import math
from datetime import datetime
from io import BytesIO
from docx import Document
from node_coordinates import NODE_COORDS


class DistanceTest:
    """
    Calculate distances between parent and child nodes in Wi-SUN network tree
    using Haversine formula based on GPS coordinates.
    """
    
    def __init__(self):
        self.coords = {k.lower(): v for k, v in NODE_COORDS.items()}
        self.ip_pattern = re.compile(r"([0-9a-fA-F:]{2,})")
        
    def parse_tree_text(self, tree_text):
        """
        Parse ASCII tree structure from text input
        Returns list of (parent, child) tuples
        """
        edges = []
        stack = []  # (depth, ip)
        
        lines = tree_text.strip().split('\n')
        for raw_line in lines:
            line = raw_line.rstrip()
            if not line.strip():
                continue
                
            match = self.ip_pattern.search(line)
            if not match:
                continue
                
            ip = match.group(1).lower()
            depth = match.start()
            
            while stack and stack[-1][0] >= depth:
                stack.pop()
                
            if stack:
                edges.append((stack[-1][1], ip))
                
            stack.append((depth, ip))
            
        return edges
    
    def haversine(self, coord1, coord2):
        """
        Calculate distance between two GPS coordinates using Haversine formula
        Returns distance in meters
        """
        lat1, lon1 = coord1
        lat2, lon2 = coord2
        R = 6371000.0  # Earth radius in meters
        
        phi1 = math.radians(lat1)
        phi2 = math.radians(lat2)
        dphi = math.radians(lat2 - lat1)
        dlambda = math.radians(lon2 - lon1)
        
        a = (math.sin(dphi / 2) ** 2 +
             math.cos(phi1) * math.cos(phi2) * math.sin(dlambda / 2) ** 2)
        
        return R * 2 * math.atan2(math.sqrt(a), math.sqrt(1 - a))
    
    def calculate_distances(self, tree_text):
        """
        Main function to calculate distances from tree text
        Returns dict with results and statistics
        """
        edges = self.parse_tree_text(tree_text)
        
        valid_rows = []
        skipped = []
        
        # Validate edge coordinates and calculate distances
        for parent, child in edges:
            if parent.lower() in self.coords and child.lower() in self.coords:
                distance = self.haversine(self.coords[parent], self.coords[child])
                valid_rows.append({
                    'parent': parent,
                    'child': child,
                    'distance': round(distance, 3)
                })
            else:
                skipped.append({
                    'parent': parent,
                    'child': child,
                    'reason': 'Missing coordinates'
                })
        
        # Sort by distance (descending)
        valid_rows_sorted = sorted(valid_rows, key=lambda x: x['distance'], reverse=True)
        
        # Calculate statistics
        stats = {}
        if valid_rows_sorted:
            distances = [row['distance'] for row in valid_rows_sorted]
            stats = {
                'total_connections': len(valid_rows_sorted),
                'max_distance': max(distances),
                'min_distance': min(distances),
                'avg_distance': sum(distances) / len(distances),
                'total_edges': len(edges),
                'skipped_edges': len(skipped)
            }
        
        return {
            'success': True,
            'data': valid_rows_sorted,
            'skipped': skipped,
            'statistics': stats
        }
    
    def generate_word_document(self, results_data):
        """
        Generate Word document with distance analysis results
        Returns BytesIO object containing the document
        """
        document = Document()
        
        # Title
        document.add_heading("Network Distance Analysis Report", level=1)
        
        # Summary paragraph
        stats = results_data.get('statistics', {})
        p = document.add_paragraph()
        p.add_run(f"Total valid connections analyzed: {stats.get('total_connections', 0)}\n")
        p.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        p.add_run("Distances calculated using Haversine formula\n")
        
        # Main table
        document.add_heading("Complete Distance Table", level=2)
        
        data = results_data.get('data', [])
        if data:
            # Create table
            table = document.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            hdr = table.rows[0].cells
            hdr[0].text = "Sr No"
            hdr[1].text = "Parent Node"
            hdr[2].text = "Child Node"
            hdr[3].text = "Distance (meters)"
            
            # Make header bold
            for cell in hdr:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Add data rows
            for i, row_data in enumerate(data, start=1):
                row = table.add_row().cells
                row[0].text = str(i)
                row[1].text = row_data['parent']
                row[2].text = row_data['child']
                row[3].text = f"{row_data['distance']:.3f}"
                
                # Add page break every 20 rows
                if i % 20 == 0 and i < len(data):
                    document.add_page_break()
                    document.add_heading(f"Distance Table (Continued) - Rows {i+1} onwards", level=2)
        
        # Statistics section
        if stats:
            document.add_page_break()
            document.add_heading("Statistics Summary", level=2)
            
            stats_table = document.add_table(rows=5, cols=2)
            stats_table.style = 'Light Shading Accent 1'
            
            stats_data = [
                ("Total Connections", str(stats.get('total_connections', 0))),
                ("Maximum Distance", f"{stats.get('max_distance', 0):.3f} m"),
                ("Minimum Distance", f"{stats.get('min_distance', 0):.3f} m"),
                ("Average Distance", f"{stats.get('avg_distance', 0):.3f} m"),
                ("Skipped Edges", str(stats.get('skipped_edges', 0)))
            ]
            
            for i, (label, value) in enumerate(stats_data):
                stats_table.rows[i].cells[0].text = label
                stats_table.rows[i].cells[1].text = value
        
        # Skipped edges section
        skipped = results_data.get('skipped', [])
        if skipped:
            document.add_page_break()
            document.add_heading("Skipped Connections (Missing Coordinates)", level=2)
            
            skip_table = document.add_table(rows=1, cols=2)
            skip_table.style = 'Light Grid Accent 1'
            
            hdr = skip_table.rows[0].cells
            hdr[0].text = "Parent Node"
            hdr[1].text = "Child Node"
            
            for cell in hdr:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            for skip_data in skipped:
                row = skip_table.add_row().cells
                row[0].text = skip_data['parent']
                row[1].text = skip_data['child']
        
        # Save to BytesIO
        doc_io = BytesIO()
        document.save(doc_io)
        doc_io.seek(0)
        
        return doc_io
